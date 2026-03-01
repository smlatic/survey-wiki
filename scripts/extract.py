#!/usr/bin/env python3
"""Extract text from Word (.docx) and Excel (.xlsx) files."""
import sys
import json
from pathlib import Path

def extract_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)

def extract_xlsx(path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True)
    output = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        output.append(f"## Sheet: {sheet_name}\n")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            output.append(" | ".join(cells))
        output.append("")
    return "\n".join(output)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract.py <file>", file=sys.stderr)
        sys.exit(1)

    filepath = sys.argv[1]
    ext = Path(filepath).suffix.lower()

    if ext == ".docx":
        print(extract_docx(filepath))
    elif ext == ".xlsx":
        print(extract_xlsx(filepath))
    else:
        print(f"Unsupported file type: {ext}", file=sys.stderr)
        sys.exit(1)
